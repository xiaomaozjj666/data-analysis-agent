from __future__ import annotations

import pandas as pd
import pytest

from data_agent.workspace import DataWorkspace


def test_workspace_loads_csv_and_profiles(workspace):
    profile = workspace.profile(sample_rows=3)
    assert profile["rows"] == 6
    assert profile["columns"] == 4
    assert profile["duplicate_rows"] == 1
    assert next(item for item in profile["column_info"] if item["name"] == "sales")["missing"] == 1


def test_workspace_reads_chinese_encoded_csv(tmp_path):
    path = tmp_path / "gb.csv"
    pd.DataFrame({"地区": ["华东"], "销售额": [100]}).to_csv(path, index=False, encoding="gb18030")
    workspace = DataWorkspace(tmp_path / "runs")
    profile = workspace.load(path)
    assert profile["rows"] == 1
    assert list(workspace.dataframe.columns) == ["地区", "销售额"]


def test_workspace_skips_unparseable_csv_rows_with_warning(tmp_path):
    path = tmp_path / "malformed.csv"
    path.write_text("a,b\n1,2\n3,4,5\n6,7\n", encoding="utf-8")
    workspace = DataWorkspace(tmp_path / "runs")

    profile = workspace.load(path)

    assert profile["rows"] == 2
    assert profile["load_warnings"]
    assert "跳过" in profile["load_warnings"][0]


def test_rejects_unsupported_file(tmp_path):
    path = tmp_path / "unsafe.py"
    path.write_text("print('no')", encoding="utf-8")
    workspace = DataWorkspace(tmp_path / "runs")
    with pytest.raises(ValueError, match="不支持"):
        workspace.load(path)


@pytest.mark.parametrize("suffix", [".xlsx", ".parquet", ".jsonl"])
def test_supported_structured_formats_round_trip(tmp_path, suffix):
    expected = pd.DataFrame({"地区": ["华东", "华南"], "销售额": [100.5, 88.0]})
    path = tmp_path / f"data{suffix}"
    if suffix == ".xlsx":
        expected.to_excel(path, index=False)
    elif suffix == ".parquet":
        expected.to_parquet(path, index=False)
    else:
        expected.to_json(path, orient="records", lines=True, force_ascii=False)
    workspace = DataWorkspace(tmp_path / "runs")
    workspace.load(path)
    assert workspace.dataframe.shape == (2, 2)
    assert workspace.dataframe["地区"].tolist() == ["华东", "华南"]


@pytest.mark.parametrize("suffix", ["csv", "xlsx", "parquet"])
def test_dataframe_export_formats(workspace, suffix):
    path = workspace.save_dataframe(f"result.{suffix}")
    assert path.exists() and path.stat().st_size > 0
    assert workspace.artifacts[-1]["kind"] == "dataset"


def test_count_artifacts_filters_by_kind(workspace):
    """count_artifacts replaces private _artifacts access from tools.py.

    The chart-numbering logic in create_visualization relies on counting
    existing visualization artifacts; this test pins the contract so a
    future refactor doesn't silently break the file-name sequence.
    """
    assert workspace.count_artifacts() == 0
    assert workspace.count_artifacts("visualization") == 0
    assert workspace.count_artifacts("dataset") == 0

    workspace.save_dataframe("result.csv")
    assert workspace.count_artifacts() == 1
    assert workspace.count_artifacts("dataset") == 1
    assert workspace.count_artifacts("visualization") == 0


def test_atomic_write_text_replaces_existing_file(tmp_path):
    """_atomic_write_text must leave a complete file even when overwriting.

    A direct write_text truncates then writes; if interrupted we'd see a
    partial file. The tmp+rename pattern guarantees the destination is
    either the old content or the new content, never a mix.
    """
    from data_agent.workspace import _atomic_write_text

    target = tmp_path / "out.html"
    target.write_text("<old>previous</old>", encoding="utf-8")
    _atomic_write_text(target, "<new>replacement</new>")
    assert target.read_text(encoding="utf-8") == "<new>replacement</new>"
    # Tmp file must be cleaned up regardless of success.
    assert not target.with_suffix(target.suffix + ".tmp").exists()


def test_atomic_write_text_cleans_up_tmp_on_failure(tmp_path):
    """If the rename fails the tmp file must not be left behind."""
    from data_agent.workspace import _atomic_write_text

    target = tmp_path / "subdir" / "missing.html"
    # Parent dir doesn't exist -> write_text raises -> tmp must be cleaned.
    with pytest.raises(OSError):
        _atomic_write_text(target, "content")
    assert not target.with_suffix(target.suffix + ".tmp").exists()


def test_profile_degrades_deep_memory_for_large_tables(tmp_path, monkeypatch):
    """超过单元格阈值时 memory_usage 降级为 deep=False，但 unique/duplicated 保持全量精确。

    把阈值 monkeypatch 为 1 强制走浅路径：profile 仍需返回 float 型 memory_mb，
    且图表语义防护依赖的 unique 计数、duplicate_rows 不受降级影响。
    """
    import data_agent.workspace as workspace_module

    monkeypatch.setattr(workspace_module, "_PROFILE_DEEP_MEMORY_MAX_CELLS", 1)
    source = tmp_path / "big.csv"
    pd.DataFrame(
        {"name": ["alpha", "beta", "alpha", "alpha"], "value": [1, 2, 2, 1]}
    ).to_csv(source, index=False)
    workspace = DataWorkspace(tmp_path / "runs")
    profile = workspace.load(source)

    assert isinstance(profile["memory_mb"], float)
    assert profile["memory_mb"] >= 0
    name_info = next(c for c in profile["column_info"] if c["name"] == "name")
    assert name_info["unique"] == 2  # 降级只影响内存统计，unique 仍全量精确
    assert profile["duplicate_rows"] == 1  # 第 4 行重复第 1 行，duplicated 仍全量精确


def test_profile_keeps_deep_memory_for_small_tables(workspace):
    """未超阈值的小表保持 deep=True 精确统计（object 列逐对象计量，值更大）。"""
    profile = workspace.profile()
    df = workspace.dataframe
    deep_mb = round(float(df.memory_usage(deep=True).sum() / 1024**2), 3)
    assert profile["memory_mb"] == deep_mb


# ---------------------------------------------------------------------------
# allocate_chart_index：图表序号原子分配
# ---------------------------------------------------------------------------


def test_allocate_chart_index_is_unique_under_concurrency(workspace):
    """回归：ToolNode 并行执行同一轮多个 create_visualization 时，
    旧实现先读 count_artifacts 再拼文件名会竞态出重号，同类型图表
    算出相同文件名后写覆盖先写，导致会话历史里图表丢失。
    新实现加锁分配，并发拿到的序号必须互不相同且连续递增。"""
    from concurrent.futures import ThreadPoolExecutor

    with ThreadPoolExecutor(max_workers=8) as pool:
        indices = list(pool.map(lambda _: workspace.allocate_chart_index(), range(32)))

    assert len(set(indices)) == 32, f"序号出现重号：{sorted(indices)}"
    assert sorted(indices) == list(range(1, 33))


def test_allocate_chart_index_resumes_after_restart(tmp_path):
    """重启/会话恢复后序号从磁盘已有图表尾号之后延续，
    不会回头覆盖历史文件；非 HTML 文件不干扰分配。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="chart_seq")
    (workspace.artifacts_dir / "折线图_3.html").write_text("<html></html>", encoding="utf-8")
    (workspace.artifacts_dir / "三维散点_7.html").write_text("<html></html>", encoding="utf-8")
    (workspace.artifacts_dir / "折线图_9.plotly.json").write_text("{}", encoding="utf-8")
    (workspace.artifacts_dir / "cleaned_data.csv").write_text("a,b\n1,2\n", encoding="utf-8")

    # 新建实例（模拟重启后恢复）：磁盘 HTML 最大尾号 7 → 下一个序号 8。
    restored = DataWorkspace(tmp_path / "runs", session_id="chart_seq")
    assert restored.allocate_chart_index() == 8
    # 高水位生效：即使序号 8 的图尚未落盘，下一次分配也不重用。
    assert restored.allocate_chart_index() == 9


# ---------------------------------------------------------------------------
# save_upload_stream：流式上传
# ---------------------------------------------------------------------------


def test_save_upload_stream_rejects_oversized_file(tmp_path):
    """超过 max_bytes 时立即报错并清理已写入的部分文件。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="stream_oversize")
    import io

    # 构造一个超过 1MB 限制的流
    payload = b"x" * (1024 * 1024 + 100)
    stream = io.BytesIO(payload)
    with pytest.raises(ValueError, match="不能超过"):
        workspace.save_upload_stream("big.csv", stream, max_bytes=1024 * 1024)
    # 部分文件必须被清理
    assert not (workspace.input_dir / "big.csv").exists()


def test_save_upload_stream_writes_normal_file(tmp_path):
    """正常流式上传应写入完整文件到 input 目录。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="stream_normal")
    import io

    content = b"region,sales\nEast,100\nWest,200\n"
    stream = io.BytesIO(content)
    path = workspace.save_upload_stream("sales.csv", stream, max_bytes=10 * 1024 * 1024)
    assert path.exists()
    assert path.read_bytes() == content


def test_save_upload_stream_rejects_unsupported_extension(tmp_path):
    """不支持的文件类型应立即报 ValueError。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="stream_bad_ext")
    import io

    with pytest.raises(ValueError, match="不支持"):
        workspace.save_upload_stream(
            "script.py", io.BytesIO(b"print(1)"), max_bytes=1024
        )


# ---------------------------------------------------------------------------
# _read_delimited_chunked：大文件分块读取
# ---------------------------------------------------------------------------


def test_read_delimited_chunked_handles_large_csv(tmp_path, monkeypatch):
    """超过 _LARGE_DELIMITED_THRESHOLD 的 CSV 走分块读取路径，结果与一次性读取一致。"""
    import data_agent.workspace as workspace_module

    # 把阈值调到 0 强制走分块路径
    monkeypatch.setattr(workspace_module, "_LARGE_DELIMITED_THRESHOLD", 0)

    df = pd.DataFrame(
        {"id": range(200), "name": [f"row_{i}" for i in range(200)], "value": [i * 1.5 for i in range(200)]}
    )
    source = tmp_path / "large.csv"
    df.to_csv(source, index=False)

    workspace = DataWorkspace(tmp_path / "runs", session_id="chunked")
    profile = workspace.load(source)

    assert profile["rows"] == 200
    assert profile["columns"] == 3
    assert workspace.dataframe["id"].tolist() == list(range(200))


def test_read_delimited_chunked_handles_bad_lines(tmp_path, monkeypatch):
    """分块读取模式下遇到坏行应回退 skip_bad=True 跳过并记录警告。"""
    import data_agent.workspace as workspace_module

    monkeypatch.setattr(workspace_module, "_LARGE_DELIMITED_THRESHOLD", 0)

    # 第二行多一列，触发 ParserError 后回退 skip
    source = tmp_path / "bad_large.csv"
    source.write_text("a,b\n1,2\n3,4,5\n6,7\n", encoding="utf-8")

    workspace = DataWorkspace(tmp_path / "runs", session_id="chunked_bad")
    profile = workspace.load(source)

    # 跳过坏行后保留 2 行有效数据
    assert profile["rows"] == 2
    assert profile["load_warnings"]


# ---------------------------------------------------------------------------
# _read_text：纯文本文件读取
# ---------------------------------------------------------------------------


def test_read_text_loads_lines_into_single_column(tmp_path):
    """TXT 文件按行解析为单列 DataFrame，保留空行。"""
    source = tmp_path / "notes.txt"
    source.write_text("第一行\n第二行\n\n第四行", encoding="utf-8")
    workspace = DataWorkspace(tmp_path / "runs", session_id="txt_normal")
    profile = workspace.load(source)

    assert profile["rows"] == 4
    assert profile["columns"] == 1
    texts = workspace.dataframe.iloc[:, 0].tolist()
    assert texts == ["第一行", "第二行", "", "第四行"]


def test_read_text_detects_gb18030_encoding(tmp_path):
    """GB18030 编码的文本文件应被正确探测并读取。"""
    source = tmp_path / "gb.txt"
    source.write_text("华东区\n华南区\n", encoding="gb18030")
    workspace = DataWorkspace(tmp_path / "runs", session_id="txt_gb")
    profile = workspace.load(source)

    assert profile["rows"] == 2
    assert workspace.dataframe.iloc[:, 0].tolist() == ["华东区", "华南区"]


def test_read_text_raises_on_empty_file(tmp_path):
    """空文本文件应报 ValueError。"""
    source = tmp_path / "empty.txt"
    source.write_text("", encoding="utf-8")
    workspace = DataWorkspace(tmp_path / "runs", session_id="txt_empty")
    with pytest.raises(ValueError, match="为空|无法解析|无法识别"):
        workspace.load(source)


# ---------------------------------------------------------------------------
# _read_pdf：PDF 表格提取（通过 mock pdfplumber）
# ---------------------------------------------------------------------------


def test_read_pdf_extracts_tables(tmp_path, monkeypatch):
    """PDF 含表格时用 pdfplumber 提取，首行做表头，多表格自动拼接。"""
    source = tmp_path / "data.pdf"
    source.write_bytes(b"%PDF-1.4 dummy")

    # 构造 mock pdfplumber，模拟两页各一个表格
    mock_page1 = type("Page", (), {"extract_tables": lambda self: [[["name", "score"], ["Alice", "90"]]]})()
    mock_page2 = type("Page", (), {"extract_tables": lambda self: [[["name", "score"], ["Bob", "85"]]]})()
    mock_pdf = type("PDF", (), {"pages": [mock_page1, mock_page2]})()

    import sys
    import types

    mock_module = types.ModuleType("pdfplumber")
    mock_module.open = lambda path: type(
        "Ctx", (), {"__enter__": lambda s: mock_pdf, "__exit__": lambda *a: None}
    )()
    monkeypatch.setitem(sys.modules, "pdfplumber", mock_module)

    workspace = DataWorkspace(tmp_path / "runs", session_id="pdf_tables")
    profile = workspace.load(source)

    assert profile["rows"] == 2
    assert workspace.dataframe["name"].tolist() == ["Alice", "Bob"]
    assert workspace.dataframe["score"].tolist() == ["90", "85"]
    assert any("PDF" in w for w in workspace.load_warnings)


def test_read_pdf_falls_back_to_text_when_no_tables(tmp_path, monkeypatch):
    """PDF 无表格时回退为按行提取文本到单列 DataFrame。"""
    source = tmp_path / "text_only.pdf"
    source.write_bytes(b"%PDF-1.4 dummy")

    mock_page = type("Page", (), {
        "extract_tables": lambda self: [],
        "extract_text": lambda self: "第一段\n第二段",
    })()
    mock_pdf = type("PDF", (), {"pages": [mock_page]})()

    import sys
    import types

    mock_module = types.ModuleType("pdfplumber")
    mock_module.open = lambda path: type(
        "Ctx", (), {"__enter__": lambda s: mock_pdf, "__exit__": lambda *a: None}
    )()
    monkeypatch.setitem(sys.modules, "pdfplumber", mock_module)

    workspace = DataWorkspace(tmp_path / "runs", session_id="pdf_text")
    profile = workspace.load(source)

    assert profile["rows"] == 2
    assert workspace.dataframe.iloc[:, 0].tolist() == ["第一段", "第二段"]


def test_read_pdf_raises_on_empty_content(tmp_path, monkeypatch):
    """PDF 无表格也无文本时应报 ValueError。"""
    source = tmp_path / "empty.pdf"
    source.write_bytes(b"%PDF-1.4 dummy")

    mock_page = type("Page", (), {
        "extract_tables": lambda self: [],
        "extract_text": lambda self: None,
    })()
    mock_pdf = type("PDF", (), {"pages": [mock_page]})()

    import sys
    import types

    mock_module = types.ModuleType("pdfplumber")
    mock_module.open = lambda path: type(
        "Ctx", (), {"__enter__": lambda s: mock_pdf, "__exit__": lambda *a: None}
    )()
    monkeypatch.setitem(sys.modules, "pdfplumber", mock_module)

    workspace = DataWorkspace(tmp_path / "runs", session_id="pdf_empty")
    with pytest.raises(ValueError, match="未找到"):
        workspace.load(source)


# ---------------------------------------------------------------------------
# _read_docx：Word 文档表格提取（使用真实 python-docx 生成文件）
# ---------------------------------------------------------------------------


def test_read_docx_extracts_tables(tmp_path):
    """DOCX 含表格时提取首行做表头，数据行拼接为 DataFrame。"""
    import docx

    source = tmp_path / "table.docx"
    doc = docx.Document()
    # 添加一个 2×3 表格
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "region"
    table.rows[0].cells[1].text = "sales"
    table.rows[1].cells[0].text = "East"
    table.rows[1].cells[1].text = "100"
    table.rows[2].cells[0].text = "West"
    table.rows[2].cells[1].text = "200"
    doc.save(str(source))

    workspace = DataWorkspace(tmp_path / "runs", session_id="docx_table")
    profile = workspace.load(source)

    assert profile["rows"] == 2
    assert workspace.dataframe["region"].tolist() == ["East", "West"]
    assert workspace.dataframe["sales"].tolist() == ["100", "200"]
    assert any("Word" in w for w in workspace.load_warnings)


def test_read_docx_falls_back_to_paragraphs(tmp_path):
    """DOCX 无表格时回退为按段落提取文本到单列 DataFrame。"""
    import docx

    source = tmp_path / "paragraphs.docx"
    doc = docx.Document()
    doc.add_paragraph("第一段内容")
    doc.add_paragraph("第二段内容")
    doc.save(str(source))

    workspace = DataWorkspace(tmp_path / "runs", session_id="docx_para")
    profile = workspace.load(source)

    assert profile["rows"] == 2
    assert workspace.dataframe.iloc[:, 0].tolist() == ["第一段内容", "第二段内容"]


def test_read_docx_raises_on_empty_document(tmp_path):
    """空 Word 文档（无表格无段落文本）应报 ValueError。"""
    import docx

    source = tmp_path / "empty.docx"
    doc = docx.Document()
    doc.save(str(source))

    workspace = DataWorkspace(tmp_path / "runs", session_id="docx_empty")
    with pytest.raises(ValueError, match="未找到"):
        workspace.load(source)


# ---------------------------------------------------------------------------
# save_dataframe：导出格式
# ---------------------------------------------------------------------------


def test_save_dataframe_rejects_unsupported_format(workspace):
    """不支持的导出格式应报 ValueError。"""
    with pytest.raises(ValueError, match="仅支持"):
        workspace.save_dataframe("result.json")


def test_save_dataframe_registers_artifact_with_description(workspace):
    """save_dataframe 注册产物时携带描述，可通过 count_artifacts 查询。"""
    before = workspace.count_artifacts("dataset")
    workspace.save_dataframe("export_with_desc.csv", description="导出的清洗结果")
    assert workspace.count_artifacts("dataset") == before + 1
    # artifacts 列表中最后一项的 description 应为传入值
    assert workspace.artifacts[-1]["description"] == "导出的清洗结果"


# ---------------------------------------------------------------------------
# save_checkpoint / restore_checkpoint：重启后恢复活动 DataFrame
# ---------------------------------------------------------------------------


def test_save_and_restore_checkpoint_round_trip(tmp_path):
    """save_checkpoint 写 parquet，restore_checkpoint 读回，数据一致。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="ckpt")
    source = tmp_path / "data.csv"
    pd.DataFrame({"a": [1, 2, 3], "b": ["x", "y", "z"]}).to_csv(source, index=False)
    workspace.load(source)

    # 修改数据后保存 checkpoint
    workspace.dataframe = pd.DataFrame({"a": [10, 20], "b": ["p", "q"]})
    ckpt_path = workspace.save_checkpoint()
    assert ckpt_path.exists()

    # 新实例（模拟重启）从 checkpoint 恢复
    restored = DataWorkspace(tmp_path / "runs", session_id="ckpt")
    assert restored.restore_checkpoint() is True
    assert restored.dataframe["a"].tolist() == [10, 20]
    assert restored.dataframe["b"].tolist() == ["p", "q"]


def test_restore_checkpoint_returns_false_when_no_state(tmp_path):
    """无 checkpoint 文件时 restore_checkpoint 返回 False。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="no_ckpt")
    assert workspace.restore_checkpoint() is False


# ---------------------------------------------------------------------------
# restore_artifacts：从目录重新注册产物文件
# ---------------------------------------------------------------------------


def test_restore_artifacts_reregisters_existing_files(tmp_path):
    """artifacts 目录中已存在的文件应被重新注册到 _artifacts。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="restore_art")
    # 在 artifacts 目录中放置文件
    chart_path = workspace.artifacts_dir / "bar_1.html"
    chart_path.write_text("<html>chart</html>", encoding="utf-8")
    data_path = workspace.artifacts_dir / "cleaned.csv"
    data_path.write_text("a,b\n1,2\n", encoding="utf-8")

    workspace.restore_artifacts()

    names = [a["name"] for a in workspace.artifacts]
    assert "bar_1.html" in names
    assert "cleaned.csv" in names
    # HTML 文件应归类为 visualization
    chart_art = next(a for a in workspace.artifacts if a["name"] == "bar_1.html")
    assert chart_art["kind"] == "visualization"
    # CSV 文件应归类为 dataset
    data_art = next(a for a in workspace.artifacts if a["name"] == "cleaned.csv")
    assert data_art["kind"] == "dataset"


def test_restore_artifacts_skips_bundle_files(tmp_path):
    """plotly.min.js 和 echarts.min.js 是共享 bundle，不应注册为产物。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="restore_skip_bundle")
    (workspace.artifacts_dir / "plotly.min.js").write_text("// plotly", encoding="utf-8")
    (workspace.artifacts_dir / "echarts.min.js").write_text("// echarts", encoding="utf-8")
    (workspace.artifacts_dir / "chart_1.html").write_text("<html></html>", encoding="utf-8")

    workspace.restore_artifacts()

    names = [a["name"] for a in workspace.artifacts]
    assert "plotly.min.js" not in names
    assert "echarts.min.js" not in names
    assert "chart_1.html" in names


def test_restore_artifacts_skips_thumbnail_cache(tmp_path):
    """*_thumb.png 是缩略图端点中间缓存，不应作为产物注册。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="restore_skip_thumb")
    (workspace.artifacts_dir / "散点图_1_thumb.png").write_bytes(b"\x89PNG\r\n\x1a\n")
    (workspace.artifacts_dir / "chart_1.html").write_text("<html></html>", encoding="utf-8")

    workspace.restore_artifacts()

    names = [a["name"] for a in workspace.artifacts]
    assert "散点图_1_thumb.png" not in names
    assert "chart_1.html" in names


def test_restore_artifacts_uses_metadata_when_provided(tmp_path):
    """传入 metadata 时应使用其中的 kind/description 而非按扩展名推断。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="restore_meta")
    (workspace.artifacts_dir / "custom.csv").write_text("a,b\n1,2\n", encoding="utf-8")

    metadata = [{"name": "custom.csv", "kind": "chart_data", "description": "自定义描述"}]
    workspace.restore_artifacts(metadata=metadata)

    art = workspace.artifacts[0]
    assert art["kind"] == "chart_data"
    assert art["description"] == "自定义描述"


# ---------------------------------------------------------------------------
# snapshot_state / restore_state：步骤级回滚
# ---------------------------------------------------------------------------


def test_snapshot_and_restore_state_rolls_back_mutations(workspace):
    """snapshot 后的数据变更和文件新增，在 restore_state 后应回滚。"""
    # 快照当前状态
    snapshot = workspace.snapshot_state()
    original_rows = len(workspace.dataframe)

    # 添加一个产物文件并修改 DataFrame
    workspace.save_dataframe("temp_export.csv")
    assert (workspace.artifacts_dir / "temp_export.csv").exists()
    assert workspace.count_artifacts() >= 1

    # 回滚
    workspace.restore_state(snapshot)
    assert len(workspace.dataframe) == original_rows
    # 新增的文件应被删除
    assert not (workspace.artifacts_dir / "temp_export.csv").exists()
    # artifacts 列表应不包含被回滚的产物
    assert all(a["name"] != "temp_export.csv" for a in workspace.artifacts)


# ---------------------------------------------------------------------------
# cleanup：清理工作区目录
# ---------------------------------------------------------------------------


def test_cleanup_removes_workspace_directory(tmp_path):
    """cleanup 应删除整个工作区目录。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="cleanup_test")
    assert workspace.root.exists()
    # 写入一些文件
    (workspace.artifacts_dir / "chart.html").write_text("<html></html>", encoding="utf-8")
    workspace.cleanup()
    assert not workspace.root.exists()


def test_cleanup_is_idempotent(tmp_path):
    """cleanup 在目录已不存在时不应报错。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="cleanup_idem")
    workspace.cleanup()
    # 再次调用不应抛出
    workspace.cleanup()


# ---------------------------------------------------------------------------
# register_artifact：产物注册校验
# ---------------------------------------------------------------------------


def test_register_artifact_rejects_file_outside_artifacts_dir(workspace, tmp_path):
    """不在 artifacts 目录内的文件不应被注册。"""
    outside = tmp_path / "outside.csv"
    outside.write_text("a,b\n1,2\n", encoding="utf-8")
    with pytest.raises(ValueError, match="artifacts 目录"):
        workspace.register_artifact(outside, "dataset", "外部文件")


# ---------------------------------------------------------------------------
# allocate_chart_index：FileNotFoundError 分支
# ---------------------------------------------------------------------------


def test_allocate_chart_index_handles_missing_artifacts_dir(tmp_path):
    """artifacts 目录被外部删除后，allocate_chart_index 不应崩溃。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="missing_dir")
    import shutil

    shutil.rmtree(workspace.artifacts_dir)
    # 目录不存在时应走 FileNotFoundError 分支，返回 1
    assert workspace.allocate_chart_index() == 1


# ---------------------------------------------------------------------------
# ensure_plotly_bundle：Plotly.js bundle 写入
# ---------------------------------------------------------------------------


def test_ensure_plotly_bundle_writes_file(tmp_path):
    """首次调用写入 plotly.min.js，再次调用复用已有文件。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="plotly_bundle")
    bundle = workspace.ensure_plotly_bundle()
    assert bundle is not None
    assert bundle.exists()
    assert bundle.name == "plotly.min.js"
    assert bundle.stat().st_size > 0

    # 第二次调用应复用已有文件（不重新写入）
    first_mtime = bundle.stat().st_mtime
    bundle2 = workspace.ensure_plotly_bundle()
    assert bundle2 == bundle
    # 文件未被重写
    assert bundle2.stat().st_mtime == first_mtime


def test_ensure_plotly_bundle_returns_none_on_import_failure(tmp_path, monkeypatch):
    """plotly 不可导入时返回 None，不阻塞调用方。"""
    import builtins

    real_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name == "plotly" or name.startswith("plotly."):
            raise ImportError("plotly not available")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    workspace = DataWorkspace(tmp_path / "runs", session_id="plotly_none")
    assert workspace.ensure_plotly_bundle() is None


# ---------------------------------------------------------------------------
# ensure_echarts_bundle：ECharts.js bundle 下载
# ---------------------------------------------------------------------------


def test_ensure_echarts_bundle_downloads_from_cdn(tmp_path, monkeypatch):
    """首次调用从 CDN 下载 echarts.min.js，后续复用。"""
    import urllib.request

    # mock urllib.request.urlopen 返回假内容
    fake_content = b"// echarts minified " + b"x" * 2048

    class FakeResponse:
        def __init__(self, content):
            self._content = content

        def read(self):
            return self._content

        def __enter__(self):
            return self

        def __exit__(self, *args):
            pass

    def fake_urlopen(url, timeout=None):
        return FakeResponse(fake_content)

    monkeypatch.setattr(urllib.request, "urlopen", fake_urlopen)

    workspace = DataWorkspace(tmp_path / "runs", session_id="echarts_bundle")
    bundle = workspace.ensure_echarts_bundle()
    assert bundle is not None
    assert bundle.exists()
    assert bundle.read_bytes() == fake_content

    # 第二次调用应复用已有文件（不重新下载）
    bundle2 = workspace.ensure_echarts_bundle()
    assert bundle2 == bundle


def test_ensure_echarts_bundle_returns_none_on_download_failure(tmp_path, monkeypatch):
    """CDN 下载失败时返回 None，调用方走 fallback。"""
    import urllib.request

    def fake_urlopen(url, timeout=None):
        raise ConnectionError("network unavailable")

    monkeypatch.setattr(urllib.request, "urlopen", fake_urlopen)

    workspace = DataWorkspace(tmp_path / "runs", session_id="echarts_fail")
    assert workspace.ensure_echarts_bundle() is None


def test_ensure_echarts_bundle_rejects_too_small_response(tmp_path, monkeypatch):
    """下载内容小于 1024 字节时视为无效，返回 None。"""
    import urllib.request

    class FakeResponse:
        def __init__(self, content):
            self._content = content

        def read(self):
            return self._content

        def __enter__(self):
            return self

        def __exit__(self, *args):
            pass

    monkeypatch.setattr(
        urllib.request,
        "urlopen",
        lambda url, timeout=None: FakeResponse(b"too small"),
    )

    workspace = DataWorkspace(tmp_path / "runs", session_id="echarts_small")
    assert workspace.ensure_echarts_bundle() is None


def test_ensure_echarts_gl_bundle_downloads_from_cdn(tmp_path, monkeypatch):
    """echarts-gl 扩展 bundle 按需下载，失败返回 None。"""
    import urllib.request

    fake_content = b"// echarts-gl minified " + b"x" * 2048

    class FakeResponse:
        def __init__(self, content):
            self._content = content

        def read(self):
            return self._content

        def __enter__(self):
            return self

        def __exit__(self, *args):
            pass

    monkeypatch.setattr(
        urllib.request,
        "urlopen",
        lambda url, timeout=None: FakeResponse(fake_content),
    )

    workspace = DataWorkspace(tmp_path / "runs", session_id="echarts_gl")
    bundle = workspace.ensure_echarts_gl_bundle()
    assert bundle is not None
    assert bundle.exists()
    assert bundle.read_bytes() == fake_content


def test_ensure_echarts_gl_bundle_returns_none_on_failure(tmp_path, monkeypatch):
    """echarts-gl 下载失败时返回 None。"""
    import urllib.request

    def fake_urlopen(url, timeout=None):
        raise ConnectionError("fail")

    monkeypatch.setattr(urllib.request, "urlopen", fake_urlopen)

    workspace = DataWorkspace(tmp_path / "runs", session_id="echarts_gl_fail")
    assert workspace.ensure_echarts_gl_bundle() is None


# ---------------------------------------------------------------------------
# restore_from_directory：通过 SessionRegistry 从磁盘目录恢复会话
# ---------------------------------------------------------------------------


def test_restore_from_directory_recovers_persisted_session(tmp_path):
    """SessionRegistry.restore_from_directory 应从磁盘目录恢复完整会话。"""
    from data_agent.registry import SessionRegistry
    from data_agent.storage import LocalSessionStorage

    runs_dir = tmp_path / "runs"
    # 先创建并持久化一个会话
    workspace = DataWorkspace(runs_dir, session_id="restore_dir_test")
    source = workspace.save_upload("sales.csv", b"region,sales\nEast,100\nWest,200\n")
    workspace.load(source)
    workspace.save_dataframe("result.csv")

    registry = SessionRegistry(runs_dir, max_sessions=10, ttl_hours=24, storage=LocalSessionStorage())
    session_id, record = registry.create(workspace)
    record.chat = [{"role": "user", "content": "检查数据"}]
    record.analysis_status = "completed"
    registry.persist(session_id, record)

    # 从内存中移除，模拟服务重启
    registry._items.clear()

    # 用 restore_from_directory 恢复
    restored = registry.restore_from_directory(session_id)
    assert restored is not None
    assert restored.workspace.dataframe.shape == (2, 2)
    assert restored.workspace.dataframe["region"].tolist() == ["East", "West"]
    assert restored.chat == [{"role": "user", "content": "检查数据"}]
    assert restored.analysis_status == "completed"
    # 产物文件应被重新注册
    assert restored.workspace.count_artifacts("dataset") >= 1


def test_restore_from_directory_returns_none_for_invalid_session(tmp_path):
    """目录无效或不存在时 restore_from_directory 返回 None。"""
    from data_agent.registry import SessionRegistry
    from data_agent.storage import LocalSessionStorage

    runs_dir = tmp_path / "runs"
    runs_dir.mkdir()
    registry = SessionRegistry(runs_dir, max_sessions=10, ttl_hours=24, storage=LocalSessionStorage())

    # 不存在的 session_id
    assert registry.restore_from_directory("nonexistent_session") is None


def test_restore_from_directory_returns_existing_when_already_in_memory(tmp_path):
    """session 已在内存中时 restore_from_directory 直接返回已有 record。"""
    from data_agent.registry import SessionRegistry
    from data_agent.storage import LocalSessionStorage

    runs_dir = tmp_path / "runs"
    workspace = DataWorkspace(runs_dir, session_id="already_loaded")
    source = workspace.save_upload("data.csv", b"a,b\n1,2\n")
    workspace.load(source)

    registry = SessionRegistry(runs_dir, max_sessions=10, ttl_hours=24, storage=LocalSessionStorage())
    session_id, record = registry.create(workspace)

    # 已在内存中，应直接返回同一 record
    restored = registry.restore_from_directory(session_id)
    assert restored is record


# ---------------------------------------------------------------------------
# 剩余分支：setter 类型校验 / 加载错误路径 / 编码探测 / PDF-DOCX 边界 / repair 分支
# ---------------------------------------------------------------------------


def test_dataframe_setter_rejects_non_dataframe(tmp_path):
    workspace = DataWorkspace(tmp_path / "runs", session_id="setter_type")
    with pytest.raises(TypeError, match="DataFrame"):
        workspace.dataframe = "not a frame"


def test_save_upload_rejects_unsupported_extension(tmp_path):
    workspace = DataWorkspace(tmp_path / "runs", session_id="upload_ext")
    with pytest.raises(ValueError, match="不支持"):
        workspace.save_upload("script.py", b"print(1)")


def test_load_raises_for_missing_file(tmp_path):
    workspace = DataWorkspace(tmp_path / "runs", session_id="load_missing")
    with pytest.raises(FileNotFoundError):
        workspace.load(tmp_path / "nonexistent.csv")


def test_load_jsonl_content_with_json_extension(tmp_path):
    """.json 扩展名但内容是 JSONL 时应回退到 lines 读取。"""
    source = tmp_path / "data.json"
    source.write_text('{"a": 1}\n{"a": 2}\n', encoding="utf-8")
    workspace = DataWorkspace(tmp_path / "runs", session_id="json_lines")
    profile = workspace.load(source)
    assert profile["rows"] == 2
    assert workspace.dataframe["a"].tolist() == [1, 2]


def test_load_rejects_empty_frame(tmp_path):
    source = tmp_path / "empty.json"
    source.write_text("[]", encoding="utf-8")
    workspace = DataWorkspace(tmp_path / "runs", session_id="load_empty")
    with pytest.raises(ValueError, match="为空|无法识别"):
        workspace.load(source)


def test_load_rejects_too_many_columns(tmp_path, monkeypatch):
    import data_agent.workspace as workspace_module

    monkeypatch.setattr(workspace_module, "_MAX_COLUMNS", 2)
    source = tmp_path / "wide.csv"
    source.write_text("a,b,c\n1,2,3\n", encoding="utf-8")
    workspace = DataWorkspace(tmp_path / "runs", session_id="load_wide")
    with pytest.raises(ValueError, match="列数超过"):
        workspace.load(source)


def test_sniff_delimiter_falls_back_to_comma(tmp_path):
    """单列 CSV 嗅探失败时应回退逗号分隔。"""
    import data_agent.workspace as workspace_module

    source = tmp_path / "single.csv"
    source.write_text("just_one_column\nvalue1\nvalue2\n", encoding="utf-8")
    assert workspace_module.DataWorkspace._sniff_delimiter(source, "utf-8") == ","


def test_read_delimited_skip_also_fails(tmp_path, monkeypatch):
    """跳过坏行后仍解析失败时应抛出明确错误（475-476 分支）。"""
    import pandas as pd

    source = tmp_path / "bad.csv"
    source.write_text("a,b\n1,2,3\n", encoding="utf-8")
    workspace = DataWorkspace(tmp_path / "runs", session_id="skip_fail")

    calls = {"n": 0}

    def flaky_read_csv(*args, **kwargs):
        calls["n"] += 1
        raise pd.errors.ParserError("simulated failure")

    monkeypatch.setattr(pd, "read_csv", flaky_read_csv)
    with pytest.raises(ValueError, match="文件格式无法解析"):
        workspace.load(source)
    assert calls["n"] >= 2  # 首次失败 + skip 重试也失败


def test_load_rejects_undecodable_csv(tmp_path):
    """探测通过但全量读取遇非法字节时应报编码错误（479-480 分支）。

    前 8192 字节必须全部合法 UTF-8（编码探测只看 8KB），非法字节置于
    探测窗口之外，read_csv 全量读取时才会触发 UnicodeDecodeError。
    """
    source = tmp_path / "binary.csv"
    source.write_bytes(b"a,b\n" + b"1,2\n" * 2050 + b"\xff\xff\xff")
    workspace = DataWorkspace(tmp_path / "runs", session_id="load_binary")
    with pytest.raises(ValueError, match="无法识别文件编码"):
        workspace.load(source)


def test_read_pdf_skips_empty_and_short_tables(tmp_path, monkeypatch):
    """PDF 中空表格与单行表格应被跳过，只提取有效表格。"""
    source = tmp_path / "mixed.pdf"
    source.write_bytes(b"%PDF-1.4 dummy")

    import sys
    import types

    mock_page = type("Page", (), {
        "extract_tables": lambda self: [
            [],  # 空表
            [["only", "header"]],  # 单行（无数据）→ 跳过
            [["name", "score"], ["Alice", "90"]],  # 有效表
        ],
        "extract_text": lambda self: None,
    })()
    mock_pdf = type("PDF", (), {"pages": [mock_page]})()
    mock_module = types.ModuleType("pdfplumber")
    mock_module.open = lambda path: type(
        "Ctx", (), {"__enter__": lambda s: mock_pdf, "__exit__": lambda *a: None}
    )()
    monkeypatch.setitem(sys.modules, "pdfplumber", mock_module)

    workspace = DataWorkspace(tmp_path / "runs", session_id="pdf_mixed")
    profile = workspace.load(source)
    assert profile["rows"] == 1
    assert workspace.dataframe["name"].tolist() == ["Alice"]


def test_read_text_truncates_long_files(tmp_path):
    source = tmp_path / "long.txt"
    source.write_text("\n".join(f"line-{i}" for i in range(10)), encoding="utf-8")
    workspace = DataWorkspace(tmp_path / "runs", session_id="txt_trunc")
    # 直接调用 _read_text 验证截断分支
    df = workspace._read_text(source, max_lines=3)
    assert len(df) == 3
    assert workspace.load_warnings and "截断" in workspace.load_warnings[0]


def test_read_text_raises_when_all_encodings_fail(tmp_path):
    source = tmp_path / "undecodable.txt"
    source.write_bytes(b"\x00\x00\xff")
    workspace = DataWorkspace(tmp_path / "runs", session_id="txt_enc")
    with pytest.raises(ValueError, match="无法识别文本文件编码"):
        workspace._read_text(source)


def test_read_docx_skips_single_row_table(tmp_path):
    """DOCX 单行表格应被跳过，回退到段落提取。"""
    import docx

    source = tmp_path / "single_row.docx"
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=2)  # 只有表头
    table.rows[0].cells[0].text = "a"
    table.rows[0].cells[1].text = "b"
    doc.add_paragraph("第一段")
    doc.save(str(source))

    workspace = DataWorkspace(tmp_path / "runs", session_id="docx_single")
    profile = workspace.load(source)
    assert profile["rows"] == 1
    assert workspace.dataframe.iloc[:, 0].tolist() == ["第一段"]


def test_repair_format_skips_empty_and_percent_columns(tmp_path):
    """repair_format 应跳过全空列、百分比/前导零列（直接注入 DataFrame 避免 CSV 推断干扰）。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="repair_edge")
    workspace.dataframe = pd.DataFrame(
        {
            "pct": ["10%", "20%"],
            "leading_zero": ["0012", "0034"],
            "all_empty": ["", ""],  # 触发缺失标记归一 → changed=True
            "note": ["x", "y"],
        }
    )
    result = workspace.repair_format(parse_dates=True)
    # 百分比列与前导零列不应被转数值（跳过分支）
    assert workspace.dataframe["pct"].dtype.kind in ("O", "U", "S")
    assert workspace.dataframe["leading_zero"].dtype.kind in ("O", "U", "S")
    assert result["changed"] is True  # "" 被归一为缺失触发导出


def test_repair_format_warns_on_unparseable_date(tmp_path):
    source = tmp_path / "bad_date.csv"
    source.write_text("date_col,value\n2025-01-01,1\nnot-a-date,2\n", encoding="utf-8")
    workspace = DataWorkspace(tmp_path / "runs", session_id="repair_date")
    workspace.load(source)
    result = workspace.repair_format()
    assert any("无法确认" in warning for warning in result["warnings"])


def test_repair_format_date_parse_fallback(tmp_path, monkeypatch):
    """format='mixed' 解析抛异常时应回退无 format 解析（680-681 分支）。"""
    import pandas as pd

    source = tmp_path / "dates.csv"
    source.write_text("date_col,value\n2025-01-01,1\n", encoding="utf-8")
    workspace = DataWorkspace(tmp_path / "runs", session_id="repair_fallback")
    workspace.load(source)

    real_to_datetime = pd.to_datetime
    calls = {"n": 0}

    def flaky_to_datetime(*args, **kwargs):
        if kwargs.get("format") == "mixed":
            calls["n"] += 1
            if calls["n"] == 1:
                raise ValueError("mixed parse error")
        return real_to_datetime(*args, **kwargs)

    monkeypatch.setattr(pd, "to_datetime", flaky_to_datetime)
    result = workspace.repair_format()
    assert pd.api.types.is_datetime64_any_dtype(workspace.dataframe["date_col"])
    assert result["changed"] is True


def test_profile_cache_lru_eviction(tmp_path, monkeypatch):
    """profile 缓存超过上限时应按 LRU 淘汰最旧条目。"""
    import data_agent.workspace as workspace_module

    monkeypatch.setattr(workspace_module, "_PROFILE_CACHE_MAX_ENTRIES", 2)
    source = tmp_path / "data.csv"
    pd.DataFrame({"a": [1, 2, 3]}).to_csv(source, index=False)
    workspace = DataWorkspace(tmp_path / "runs", session_id="cache_lru")
    workspace.load(source)
    for sample in (1, 2, 3, 4, 5):
        workspace.profile(sample_rows=sample)
    assert len(workspace._profile_cache) <= 2


def test_ensure_echarts_gl_bundle_reuses_existing_file(tmp_path, monkeypatch):
    """echarts-gl bundle 已存在时应直接复用，不触发下载。"""
    import urllib.request

    def unexpected(*args, **kwargs):
        raise AssertionError("不应发起网络下载")

    monkeypatch.setattr(urllib.request, "urlopen", unexpected)
    workspace = DataWorkspace(tmp_path / "runs", session_id="gl_reuse")
    bundle = workspace.artifacts_dir / "echarts-gl.min.js"
    bundle.parent.mkdir(parents=True, exist_ok=True)
    bundle.write_text("/* existing gl */", encoding="utf-8")
    assert workspace.ensure_echarts_gl_bundle() == bundle


def test_ensure_echarts_gl_bundle_rejects_small_response(tmp_path, monkeypatch):
    """echarts-gl 下载内容过小时视为无效返回 None。"""
    import urllib.request

    class FakeResponse:
        def read(self):
            return b"too small"

        def __enter__(self):
            return self

        def __exit__(self, *args):
            pass

    monkeypatch.setattr(
        urllib.request, "urlopen", lambda url, timeout=None: FakeResponse()
    )
    workspace = DataWorkspace(tmp_path / "runs", session_id="gl_small")
    assert workspace.ensure_echarts_gl_bundle() is None


def test_snapshot_state_handles_missing_artifacts_dir(tmp_path):
    """artifacts 目录被删除后 snapshot_state 不应崩溃。"""
    import shutil

    workspace = DataWorkspace(tmp_path / "runs", session_id="snap_missing")
    workspace.dataframe = pd.DataFrame({"a": [1, 2]})
    shutil.rmtree(workspace.artifacts_dir)
    df, files, version = workspace.snapshot_state()
    assert files == set()
    assert df is not None


def test_restore_state_restores_dataframe_and_handles_missing_dir(tmp_path):
    """restore_state 应恢复被修改的 DataFrame 并容忍 artifacts 目录缺失。"""
    import shutil

    workspace = DataWorkspace(tmp_path / "runs", session_id="restore_missing")
    workspace.dataframe = pd.DataFrame({"a": [1, 2]})
    snapshot = workspace.snapshot_state()
    workspace.dataframe = pd.DataFrame({"a": [99]})  # 修改数据
    shutil.rmtree(workspace.artifacts_dir)  # 删除产物目录
    workspace.restore_state(snapshot)
    assert workspace.dataframe["a"].tolist() == [1, 2]


def test_restore_artifacts_returns_early_without_dir(tmp_path):
    import shutil

    workspace = DataWorkspace(tmp_path / "runs", session_id="restore_no_dir")
    # 删除 artifacts 目录（DataWorkspace 初始化时会创建），触发 948 早退分支
    shutil.rmtree(workspace.artifacts_dir)
    workspace.restore_artifacts()
    assert workspace.artifacts == []


def test_sniff_encoding_falls_back_to_utf8(tmp_path):
    """所有候选编码都失败时 _sniff_encoding 回退 utf-8（411 分支）。"""
    source = tmp_path / "weird.bin"
    source.write_bytes(b"\x00\x00\xff")
    assert DataWorkspace._sniff_encoding(source) == "utf-8"


def test_repair_format_skips_empty_date_column(tmp_path):
    """日期列全为空时跳过解析（683 分支）。"""
    workspace = DataWorkspace(tmp_path / "runs", session_id="repair_empty_date")
    workspace.dataframe = pd.DataFrame({"date_col": [None, None], "v": [1, 2]})
    result = workspace.repair_format()
    assert result["changed"] is False
