"""数据工作区管理：为每个分析会话提供隔离的文件系统和 DataFrame 生命周期管理。

核心职责：
- 文件加载：支持 CSV/TSV/Excel/JSON/JSONL/Parquet，自动探测编码和分隔符。
- 产物管理：图表、清洗数据、导出文件统一注册在 artifacts 目录。
- 状态快照/回滚：每步执行前 snapshot，失败时 restore，保证数据一致性。
- 持久化：checkpoint 保存活动 DataFrame，重启后可恢复。
- Profile 缓存：避免 finalize 等节点重复计算开销显著的概况统计。

线程安全：
    DataWorkspace 不是线程安全的。API 层通过 SessionRecord.run_lock
    保证同一会话同一时刻只有一个分析线程在操作 workspace。
    例外：同一分析线程内 ToolNode 会并行执行同一轮的多个工具调用，
    图表序号分配因此必须原子化（见 ``allocate_chart_index``），否则并行
    生成的同类型图表会算出相同文件名互相覆盖。
"""

from __future__ import annotations

import csv as _csv
import logging
import os
import re
import shutil
import threading
from collections import OrderedDict
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from uuid import uuid4

import pandas as pd

from data_agent.serialization import to_jsonable

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 命名常量
# ---------------------------------------------------------------------------

#: 支持的数据文件扩展名集合。
#: 结构化表格：CSV/TSV/Excel/JSON/JSONL/Parquet
#: 非结构化/半结构化：PDF（表格提取）、TXT（行式文本）、DOCX（Word 表格）
SUPPORTED_EXTENSIONS = {
    ".csv", ".tsv", ".xlsx", ".xls", ".json", ".jsonl", ".parquet",
    ".pdf", ".txt", ".docx",
}

#: 共享 Plotly.js 束缚文件名，每个工作区只写一次，所有图表复用。
PLOTLY_BUNDLE_NAME = "plotly.min.js"
#: ECharts 引擎所需的前端 bundle 文件名，与 plotly.min.js 同目录共存。
#: 双引擎互不冲突：HTML 通过相对路径引用各自的 bundle。
ECHARTS_BUNDLE_NAME = "echarts.min.js"
#: ECharts 官方稳定版 CDN，首次生成 echarts 图表时下载到 artifacts_dir，
#: 后续复用。下载失败时 fallback 到 CDN URL 直接引用（在线场景）。
ECHARTS_CDN_URL = "https://cdn.jsdelivr.net/npm/echarts@5.5.0/dist/echarts.min.js"
#: echarts-gl 扩展 bundle（3D 散点等 gl 系列需要），与主 bundle 同目录按需下载。
ECHARTS_GL_BUNDLE_NAME = "echarts-gl.min.js"
ECHARTS_GL_CDN_URL = "https://cdn.jsdelivr.net/npm/echarts-gl@2.0.9/dist/echarts-gl.min.js"

#: 活动 DataFrame 的 checkpoint 文件名，用于重启后恢复。
WORKSPACE_STATE_NAME = "workspace_state.parquet"

#: 流式上传时的分块大小（1 MB），平衡内存占用和 I/O 次数。
_UPLOAD_CHUNK_SIZE = 1024 * 1024

#: 拒绝加载的最大列数，防止意外资源耗尽。
_MAX_COLUMNS = 1000

#: 超过此大小（字节）的 CSV/TSV 使用分块读取，避免 pandas 解析时
#: 在内存中构建中间数据结构导致瞬时内存尖峰（可达文件大小的 3-5 倍）。
#: 分块读取用 TextFileReader 流式解析，每块读入后立即 append，峰值内存
#: 降为约 1.5 倍最终 DataFrame 大小。50MB 阈值覆盖大多数业务 CSV。
_LARGE_DELIMITED_THRESHOLD = 50 * 1024 * 1024

#: 分块读取的块大小（行数）。太小则 I/O 次数多，太大则峰值内存高。
#: 5 万行/块在 10 列场景下约 5-10MB/块，平衡两者。
_CHUNK_SIZE = 50_000

#: Profile 缓存最大条目数，超过后按 LRU 策略淘汰最旧条目。
#: 值为 4：覆盖常见场景（validate_dataset 用 sample_rows=5，finalize 用
#: sample_rows=5，inspect_data 工具可能用不同 sample_rows），留余量避免
#: 频繁淘汰导致缓存命中率下降。
_PROFILE_CACHE_MAX_ENTRIES = 4

#: CSV 编码探测顺序：UTF-8 BOM → UTF-8 → GB18030（覆盖中文 Windows 场景）。
_CSV_ENCODING_CANDIDATES = ("utf-8-sig", "utf-8", "gb18030")

#: profile() 中深度内存统计（memory_usage(deep=True)）的单元格上限。
#: deep=True 需要逐个遍历 Python 对象（object 列每个字符串都要 sizeof），
#: 大表（如 50 万行 × 30 列）会卡住上传接口数秒生成首次 profile。
#: 超过阈值时降级为 deep=False 估算（object 列按指针大小计，偏小但秒回）；
#: nunique / duplicated 保持全量——它们是 C 层哈希实现且图表语义防护
#: （标识符列/常量列判定）依赖精确的 unique 计数，不能采样。
_PROFILE_DEEP_MEMORY_MAX_CELLS = 5_000_000


def _atomic_write_text(path: Path, content: str, *, encoding: str = "utf-8") -> None:
    """Write text atomically: write to a sibling .tmp file then rename.

    A direct ``path.write_text`` truncates the destination before writing; if
    the process is killed mid-write (OOM, deploy restart, disk full) we leave
    a corrupt partial file that subsequent reads will fail on. The tmp + rename
    pattern guarantees readers either see the old file or the new file, never
    a half-written one. ``os.replace`` is atomic on POSIX and Windows for
    same-filesystem renames.
    """
    temporary = path.with_suffix(path.suffix + ".tmp")
    try:
        temporary.write_text(content, encoding=encoding)
        os.replace(temporary, path)
    except Exception:
        temporary.unlink(missing_ok=True)
        raise


def _downcast_dtypes(df: pd.DataFrame) -> pd.DataFrame:
    """降级 DataFrame 数值列的 dtype 以节省内存。

    pandas 默认用 int64/float64，对大多数业务数据来说 int32/float32 已足够。
    安全降级：仅在不丢失精度时转换（如 0-255 的整数列 → uint8）。
    典型场景：100 万行 × 10 列的 int64 数据，降级后内存从 ~80MB 降到 ~20MB。
    """
    for col in df.columns:
        col_data = df[col]
        if col_data.dtype == "int64":
            # 尝试降级到最小可容纳的整数类型
            df[col] = pd.to_numeric(col_data, downcast="integer")
        elif col_data.dtype == "float64":
            # float32 精度足够大多数统计场景（6-7 位有效数字）
            df[col] = pd.to_numeric(col_data, downcast="float")
    return df


@dataclass(frozen=True, slots=True)
class Artifact:
    """工作区中已注册的产物文件元数据。

    Attributes:
        name: 文件名（不含目录）。
        kind: 产物类型（visualization / dataset / image / chart_data）。
        path: 绝对路径。
        description: 面向用户的简短描述。
    """

    name: str
    kind: str
    path: Path
    description: str

    def as_dict(self) -> dict[str, str]:
        return {
            "name": self.name,
            "kind": self.kind,
            "path": str(self.path.resolve()),
            "description": self.description,
        }


class DataWorkspace:
    """单个分析会话的数据工作区，管理活动 DataFrame 和生成的产物文件。

    目录结构::

        <root>/<session_id>/
        ├── input/          # 上传的原始数据文件
        ├── artifacts/      # 图表、清洗数据、导出文件等产物
        └── workspace_state.parquet  # 活动 DataFrame 的 checkpoint

    生命周期：
        1. 构造时创建目录结构。
        2. load() 加载数据文件到内存。
        3. Agent 工具通过 dataframe 属性读写数据。
        4. save_checkpoint() / restore_checkpoint() 支持重启恢复。
        5. cleanup() 在会话过期时删除整个目录。
    """

    def __init__(self, root: str | Path, session_id: str | None = None) -> None:
        safe_id = re.sub(r"[^a-zA-Z0-9_-]", "_", session_id or uuid4().hex)
        self.root = Path(root).expanduser().resolve() / safe_id
        self.input_dir = self.root / "input"
        self.artifacts_dir = self.root / "artifacts"
        self.input_dir.mkdir(parents=True, exist_ok=True)
        self.artifacts_dir.mkdir(parents=True, exist_ok=True)
        self._df: pd.DataFrame | None = None
        self._source_row_count = 0
        self.source_path: Path | None = None
        self._artifacts: list[Artifact] = []
        self.load_warnings: list[str] = []
        # Profile 缓存：避免 finalize 等节点重复计算 profile（对大 DataFrame
        # 的 duplicated().sum() 和 nunique() 开销显著）。数据变更时自动失效。
        # 用 OrderedDict 实现 LRU：命中时 move_to_end，超容量时 popitem(last=False)。
        self._profile_cache: OrderedDict[int, dict[str, Any]] = OrderedDict()
        # 全量统计缓存：nunique()、duplicated().sum()、memory_usage() 等昂贵
        # 全表扫描独立于 sample_rows，单独缓存一份供不同 sample_rows 的
        # profile() 调用复用，避免 upload（sample_rows=5）和 GET session
        # （sample_rows=8）各扫一遍全表。
        self._cached_full_stats: dict[str, Any] | None = None
        self._cached_full_stats_version: int = -1
        self._df_version = 0  # 每次 setter 递增，避免 id() 复用导致过期缓存
        # 图表序号分配：ToolNode 并行执行同一轮多个 create_visualization 时，
        # 先读 count_artifacts 再拼文件名会竞态出重号（同类型图相互覆盖）。
        # 用锁 + 进程内高水位保证序号单调递增；磁盘扫描兼顾重启后延续。
        self._chart_index_lock = threading.Lock()
        self._chart_index_high_water = 0

    @property
    def dataframe(self) -> pd.DataFrame:
        if self._df is None:
            raise RuntimeError("尚未加载数据集。")
        return self._df

    @dataframe.setter
    def dataframe(self, value: pd.DataFrame) -> None:
        if not isinstance(value, pd.DataFrame):
            raise TypeError("工作区数据必须是 pandas DataFrame。")
        self._df = value
        self._df_version += 1
        # 数据变更时清除 profile 缓存，确保下次 profile() 反映最新数据。
        self._profile_cache.clear()
        self._cached_full_stats = None

    @property
    def artifacts(self) -> list[dict[str, str]]:
        return [item.as_dict() for item in self._artifacts]

    def count_artifacts(self, kind: str | None = None) -> int:
        """Count registered artifacts, optionally filtered by kind.

        Public accessor so tools and API code don't need to reach into the
        private ``_artifacts`` list. Filtering by kind lets callers answer
        "how many charts exist?" without materialising the full dict list.
        """
        if kind is None:
            return len(self._artifacts)
        return sum(1 for item in self._artifacts if item.kind == kind)

    def allocate_chart_index(self) -> int:
        """线程安全地分配全局递增的图表序号（从 1 开始）。

        首次调用时会扫描 artifacts 目录中现有 HTML 文件的尾号以校准
        高水位（支持服务重启后序号延续），后续分配仅在进程内递增计数器，
        不再重复扫描目录——即 O(1) 而非 O(n)。
        """
        with self._chart_index_lock:
            if self._chart_index_high_water == 0:
                # First allocation after init or restore: seed from disk.
                highest = 0
                try:
                    children = list(self.artifacts_dir.iterdir())
                except FileNotFoundError:
                    children = []
                for path in children:
                    if path.suffix.lower() != ".html":
                        continue
                    match = re.search(r"_(\d+)$", path.stem)
                    if match:
                        highest = max(highest, int(match.group(1)))
                self._chart_index_high_water = highest
            self._chart_index_high_water += 1
            return self._chart_index_high_water

    def register_artifact(self, path: str | Path, kind: str, description: str) -> Artifact:
        resolved = Path(path).resolve()
        if self.artifacts_dir.resolve() not in resolved.parents:
            raise ValueError("产物文件必须位于当前工作区 artifacts 目录内。")
        artifact = Artifact(resolved.name, kind, resolved, description)
        if all(existing.path != resolved for existing in self._artifacts):
            self._artifacts.append(artifact)
        return artifact

    def save_upload(self, filename: str, content: bytes) -> Path:
        safe_name = re.sub(r"[^\w.\-()\u4e00-\u9fff]", "_", Path(filename).name)
        if Path(safe_name).suffix.lower() not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件类型。支持：{', '.join(sorted(SUPPORTED_EXTENSIONS))}")
        destination = (self.input_dir / safe_name).resolve()
        destination.write_bytes(content)
        return destination

    def save_upload_stream(self, filename: str, stream: Any, max_bytes: int) -> Path:
        """流式写入上传文件，避免在内存中复制整个文件。

        以 _UPLOAD_CHUNK_SIZE 分块读取，超过 max_bytes 时立即报错并清理
        已写入的部分文件。

        Args:
            filename: 原始文件名。
            stream: 文件类对象（支持 read(size)）。
            max_bytes: 允许的最大字节数。

        Returns:
            写入的目标文件路径。

        Raises:
            ValueError: 文件类型不支持或超过大小限制。
        """
        safe_name = re.sub(r"[^\w.\-()\u4e00-\u9fff]", "_", Path(filename).name)
        if Path(safe_name).suffix.lower() not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件类型。支持：{', '.join(sorted(SUPPORTED_EXTENSIONS))}")
        destination = (self.input_dir / safe_name).resolve()
        total = 0
        try:
            with destination.open("wb") as target:
                while True:
                    chunk = stream.read(_UPLOAD_CHUNK_SIZE)
                    if not chunk:
                        break
                    total += len(chunk)
                    if total > max_bytes:
                        raise ValueError(f"文件不能超过 {max_bytes // (1024 * 1024)}MB。")
                    target.write(chunk)
        except Exception:
            destination.unlink(missing_ok=True)
            raise
        return destination

    def load(self, source: str | Path, copy_into_workspace: bool = False) -> dict[str, Any]:
        """加载数据文件到工作区并返回数据概况。

        支持结构化表格（CSV/TSV/Excel/Parquet/JSON/JSONL）与非结构化数据
        （PDF 表格提取、TXT 行式文本、DOCX Word 表格）。
        加载失败时抛出 ValueError 或 FileNotFoundError，不会留下部分状态。

        Args:
            source: 数据文件路径。
            copy_into_workspace: 是否将文件复制到 input/ 目录（部署场景用）。

        Returns:
            数据概况字典（同 profile() 返回值）。

        Raises:
            FileNotFoundError: 文件不存在。
            ValueError: 格式不支持、解析失败、列数超限。
        """
        path = Path(source).expanduser().resolve()
        if not path.is_file():
            raise FileNotFoundError(f"数据文件不存在：{path}")
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持 {suffix} 文件。支持：{', '.join(sorted(SUPPORTED_EXTENSIONS))}")
        if copy_into_workspace and self.input_dir.resolve() not in path.parents:
            target = self.input_dir / path.name
            shutil.copy2(path, target)
            path = target.resolve()

        self.load_warnings = []
        try:
            if suffix in {".csv", ".tsv"}:
                df = self._read_delimited(path, suffix)
            elif suffix in {".xlsx", ".xls"}:
                df = pd.read_excel(path)
            elif suffix == ".parquet":
                df = pd.read_parquet(path)
            elif suffix == ".jsonl":
                df = pd.read_json(path, lines=True)
            elif suffix == ".json":
                try:
                    df = pd.read_json(path)
                except ValueError:
                    df = pd.read_json(path, lines=True)
            elif suffix == ".pdf":
                df = self._read_pdf(path)
            elif suffix == ".txt":
                df = self._read_text(path)
            elif suffix == ".docx":
                df = self._read_docx(path)
            else:  # pragma: no cover - 已被 SUPPORTED_EXTENSIONS 守护
                raise ValueError(f"未实现的格式：{suffix}")
        except (pd.errors.ParserError, UnicodeDecodeError, ValueError) as exc:
            raise ValueError(f"文件格式无法解析：{exc}") from exc

        if df.empty and len(df.columns) == 0:
            raise ValueError("数据文件为空或无法识别出列。")
        if len(df.columns) > _MAX_COLUMNS:
            raise ValueError(f"列数超过 {_MAX_COLUMNS}，拒绝加载以防止意外资源耗尽。")
        # 通过 setter 赋值，确保 _profile_cache 被清除（虽然 load 是首次设置，
        # 缓存此时为空，但走 setter 保持一致性，避免未来重构引入缓存失效 bug）。
        self.dataframe = df
        self._source_row_count = len(df)
        self.source_path = path
        return self.profile(sample_rows=5)

    @staticmethod
    def _sniff_encoding(path: Path) -> str:
        """Read the first 8 KB to determine the file encoding.

        Returns the first encoding from ``_CSV_ENCODING_CANDIDATES`` that
        successfully decodes the sample, or ``'utf-8'`` as a fallback.
        """
        with open(path, "rb") as fh:
            raw = fh.read(8192)
        for enc in _CSV_ENCODING_CANDIDATES:
            try:
                raw.decode(enc)
                return enc
            except UnicodeDecodeError:
                continue
        return "utf-8"

    @staticmethod
    def _sniff_delimiter(path: Path, encoding: str) -> str:
        """Sniff CSV delimiter from the first 8 KB of text.

        Uses ``csv.Sniffer`` to detect the separator. Falls back to ``','``
        when sniffing fails (single-column CSV, empty files, etc.).
        """
        with open(path, encoding=encoding) as fh:
            sample = fh.read(8192)
        try:
            dialect = _csv.Sniffer().sniff(sample, delimiters=",\t|;")
            return dialect.delimiter
        except _csv.Error:
            return ","

    def _read_delimited(self, path: Path, suffix: str) -> pd.DataFrame:
        """读取分隔符文件，嗅探编码和分隔符后一次解析。

        - 从文件前 8 KB 嗅探编码，避免按 3 种候选编码各读一遍文件。
        - CSV 文件（suffix 不是 .tsv）从首 8 KB 嗅探分隔符，直接用 C 引擎
          （速度约为 Python 引擎的 3-5 倍）解析，避免 sep=None 的 auto-detect
          走 Python 引擎。
        - 大文件（超 ``_LARGE_DELIMITED_THRESHOLD``）使用分块流式读取。
        - 首次尝试严格模式（on_bad_lines='error'），失败后回退到跳过坏行
          并记录警告。
        """
        encoding = self._sniff_encoding(path)
        if suffix == ".tsv":
            sep: str = "\t"
        else:
            sep = self._sniff_delimiter(path, encoding)
        file_size = path.stat().st_size
        use_chunks = file_size > _LARGE_DELIMITED_THRESHOLD
        kwargs: dict[str, Any] = {"encoding": encoding, "on_bad_lines": "error"}
        # Use C engine with explicit sep — it is 3-5× faster than the Python
        # engine used by sep=None auto-detection.
        if not use_chunks:
            kwargs["engine"] = "c"
            kwargs["sep"] = sep
        try:
            if use_chunks:
                reader_kwargs = {**kwargs, "chunksize": _CHUNK_SIZE}
                chunks: list[pd.DataFrame] = []
                for chunk in pd.read_csv(path, sep=sep, engine="python", **reader_kwargs):  # type: ignore[arg-type]
                    chunks.append(chunk)
                df: pd.DataFrame = pd.concat(chunks, ignore_index=True) if chunks else pd.DataFrame()
            else:
                df = pd.read_csv(path, **kwargs)  # type: ignore[arg-type]
            return _downcast_dtypes(df)
        except pd.errors.ParserError:
            try:
                if use_chunks:
                    reader_kwargs["on_bad_lines"] = "skip"
                    chunks = []
                    for chunk in pd.read_csv(path, sep=sep, engine="python", **reader_kwargs):  # type: ignore[arg-type]
                        chunks.append(chunk)
                    repaired = pd.concat(chunks, ignore_index=True) if chunks else pd.DataFrame()
                else:
                    kwargs["on_bad_lines"] = "skip"
                    kwargs["engine"] = "c"
                    repaired = pd.read_csv(path, **kwargs)  # type: ignore[arg-type]
                repaired = _downcast_dtypes(repaired)
            except pd.errors.ParserError:
                raise ValueError("文件格式无法解析，请检查数据格式。") from None
            self.load_warnings.append("文件包含格式异常行，已跳过无法解析的记录。")
            return repaired
        except UnicodeDecodeError:
            raise ValueError(
                "无法识别文件编码：已尝试 UTF-8、GB18030 等常见编码均失败。"
                "请用 Excel 或文本编辑器将文件另存为 UTF-8 编码的 CSV 后重新上传。"
            ) from None

    def _read_pdf(self, path: Path) -> pd.DataFrame:
        """从 PDF 文件提取表格数据。

        使用 pdfplumber 逐页提取表格，多页表格自动拼接。
        若 PDF 中无表格（纯文本 PDF），回退为按行提取文本到单列 DataFrame。
        """
        import pdfplumber

        tables: list[pd.DataFrame] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_tables = page.extract_tables()
                for table in page_tables:
                    if not table or len(table) < 2:
                        continue
                    # 首行作为表头，其余为数据行
                    header = [str(c or f"列{i+1}").strip() for i, c in enumerate(table[0])]
                    data = table[1:]
                    df_page = pd.DataFrame(data, columns=header)
                    tables.append(df_page)

        if tables:
            # 多表格拼接：列名对齐取并集，缺失列填 NaN
            df = pd.concat(tables, ignore_index=True)
            self.load_warnings.append(f"从 PDF 提取了 {len(tables)} 个表格，共 {len(df)} 行。")
            return df

        # 回退：PDF 无表格，按行提取纯文本
        lines: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    lines.extend(text.splitlines())
        if not lines:
            raise ValueError("PDF 文件中未找到表格或文本内容。")
        self.load_warnings.append("PDF 中未检测到表格，已按行提取文本内容。")
        return pd.DataFrame({"文本": lines})

    def _read_text(self, path: Path, max_lines: int = 500_000) -> pd.DataFrame:
        """读取纯文本文件，按行解析为单列 DataFrame。

        自动探测编码（与 CSV 相同的候选列表）。超长文件截断到 max_lines 行
        并记录警告，防止日志类文件撑爆内存。
        """
        last_error: Exception | None = None
        for encoding in _CSV_ENCODING_CANDIDATES:
            try:
                with open(path, encoding=encoding) as f:
                    lines = f.readlines(max_lines + 1)
                truncated = len(lines) > max_lines
                if truncated:
                    lines = lines[:max_lines]
                    self.load_warnings.append(
                        f"文本文件超过 {max_lines} 行，已截断。请使用更专业的日志分析工具处理超大文件。"
                    )
                # 去除行尾换行符，空行保留为空字符串
                content = [line.rstrip("\n\r") for line in lines]
                if not content:
                    raise ValueError("文本文件为空。")
                return pd.DataFrame({"文本": content})
            except UnicodeDecodeError as exc:
                last_error = exc
        raise ValueError(
            "无法识别文本文件编码：已尝试 UTF-8、GB18030 等常见编码均失败。"
            f"（技术详情：{last_error}）"
        )

    def _read_docx(self, path: Path) -> pd.DataFrame:
        """从 Word 文档提取表格数据。

        使用 python-docx 读取文档中的所有表格，多表格自动拼接。
        若文档中无表格，回退为按段落提取文本到单列 DataFrame。
        """
        import docx

        doc = docx.Document(str(path))
        tables: list[pd.DataFrame] = []

        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            # 首行作为表头
            header = [cell.text.strip() or f"列{i+1}" for i, cell in enumerate(table.rows[0].cells)]
            data = [[cell.text.strip() for cell in row.cells] for row in table.rows[1:]]
            tables.append(pd.DataFrame(data, columns=header))

        if tables:
            df = pd.concat(tables, ignore_index=True)
            self.load_warnings.append(f"从 Word 文档提取了 {len(tables)} 个表格，共 {len(df)} 行。")
            return df

        # 回退：文档无表格，按段落提取文本
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            raise ValueError("Word 文档中未找到表格或文本内容。")
        self.load_warnings.append("Word 文档中未检测到表格，已按段落提取文本内容。")
        return pd.DataFrame({"文本": paragraphs})

    def repair_format(
        self,
        *,
        normalize_missing: bool = True,
        trim_strings: bool = True,
        parse_numeric: bool = True,
        parse_dates: bool = True,
        normalize_column_names: bool = False,
    ) -> dict[str, Any]:
        """Apply conservative, auditable repairs for unambiguous formatting issues."""
        df = self.dataframe.copy()
        before = {
            "rows": len(df),
            "columns": len(df.columns),
            "missing": int(df.isna().sum().sum()),
        }
        changes: list[str] = []
        warnings: list[str] = []

        if normalize_column_names:
            original = list(df.columns)
            normalized: list[str] = []
            seen: dict[str, int] = {}
            for value in original:
                base = re.sub(r"[^\w\u4e00-\u9fff]+", "_", str(value).strip().lower()).strip("_") or "column"
                seen[base] = seen.get(base, 0) + 1
                normalized.append(base if seen[base] == 1 else f"{base}_{seen[base]}")
            if normalized != original:
                df.columns = normalized
                changes.append("规范化列名并处理重复列名")

        if trim_strings:
            string_columns = list(df.select_dtypes(include=["object", "string"]).columns)
            trimmed = 0
            for column in string_columns:
                original_values = df[column].copy()
                df[column] = df[column].map(lambda value: value.strip() if isinstance(value, str) else value)
                if not df[column].equals(original_values):
                    trimmed += 1
            if trimmed:
                changes.append(f"清理 {trimmed} 个文本列首尾空格")

        if normalize_missing:
            missing_tokens = {"", "na", "n/a", "null", "none", "nan", "<na>"}
            normalized_missing = 0
            for column in df.select_dtypes(include=["object", "string"]).columns:
                values = df[column].astype("string")
                mask = values.str.strip().str.lower().isin(missing_tokens)
                normalized_missing += int(mask.sum())
                df.loc[mask, column] = pd.NA
            if normalized_missing:
                changes.append(f"将 {normalized_missing} 个明确缺失标记统一为空值")

        if parse_numeric:
            # Skip columns that carry semantic suffixes (units like kg, 元,
            # °C, km/h, ...) since converting them would silently drop the
            # unit and corrupt the data. Only currency prefixes and thousands
            # separators are considered unambiguous and get stripped.
            unit_residual_pattern = re.compile(
                r"^[\s¥￥$€]*[+-]?[\d.,]+(?:[eE][+-]?\d+)?"
            )
            for column in list(df.select_dtypes(include=["object", "string"]).columns):
                series = df[column]
                non_empty = series.dropna()
                if non_empty.empty:
                    continue
                text = non_empty.astype("string").str.strip()
                if text.str.contains("%", regex=False).any() or text.str.match(r"^0\d+$").any():
                    continue
                stripped = text.str.replace(",", "", regex=False).str.replace(
                    r"^[¥￥$€]\s*", "", regex=True
                )
                # After removing the leading number, anything non-empty that
                # remains is a unit suffix -> skip the column.
                residual = stripped.str.replace(unit_residual_pattern, "", regex=True).str.strip()
                if (residual.str.len() > 0).any():
                    continue
                candidate = pd.to_numeric(stripped, errors="coerce")
                if candidate.notna().all():
                    df[column] = pd.to_numeric(
                        df[column].astype("string").str.strip()
                        .str.replace(",", "", regex=False)
                        .str.replace(r"^[¥￥$€]\s*", "", regex=True),
                        errors="coerce",
                    )
                    changes.append(f"将格式明确的数值列 {column} 转为数值类型")

        if parse_dates:
            for column in list(df.select_dtypes(include=["object", "string"]).columns):
                if not re.search(r"date|time|日期|时间", str(column), flags=re.IGNORECASE):
                    continue
                non_empty = df[column].dropna()
                if non_empty.empty:
                    continue
                try:
                    parsed = pd.to_datetime(non_empty, errors="coerce", format="mixed")
                except (TypeError, ValueError):
                    parsed = pd.to_datetime(non_empty, errors="coerce")
                if parsed.notna().all():
                    df[column] = pd.to_datetime(df[column], errors="coerce", format="mixed")
                    changes.append(f"将日期列 {column} 转为日期时间类型")
                else:
                    warnings.append(f"日期列 {column} 存在无法确认的值，未自动转换。")

        changed = bool(changes)
        output: Path | None = None
        if changed:
            self.dataframe = df.reset_index(drop=True)
            output = self.save_dataframe("format_repaired.csv")
        after = {
            "rows": len(self.dataframe),
            "columns": len(self.dataframe.columns),
            "missing": int(self.dataframe.isna().sum().sum()),
        }
        return {
            "status": "ok",
            "changed": changed,
            "before": before,
            "after": after,
            "changes": changes,
            "warnings": warnings,
            "output": str(output) if output else None,
        }

    def _compute_full_stats(self) -> dict[str, Any]:
        """Compute expensive full-table stats once and cache across sample_rows.

        ``nunique()`` and ``duplicated().sum()`` both do full-table scans.
        When ``profile(sample_rows=5)`` and ``profile(sample_rows=8)`` are
        called in succession (upload + session GET), caching these avoids
        scanning the table twice.
        """
        if self._cached_full_stats is not None and self._cached_full_stats_version == self._df_version:
            return self._cached_full_stats
        df = self.dataframe
        missing = df.isna().sum()
        unique = df.nunique(dropna=True)
        deep_memory = len(df) * max(len(df.columns), 1) <= _PROFILE_DEEP_MEMORY_MAX_CELLS
        missing_values = list(missing)
        unique_values = list(unique)
        dtype_values = list(df.dtypes)
        column_info = [
            {
                "name": str(column),
                "dtype": str(dtype_values[i]),
                "missing": int(missing_values[i]),
                "missing_pct": round(float(missing_values[i] / max(len(df), 1) * 100), 2),
                "unique": int(unique_values[i]),
            }
            for i, column in enumerate(df.columns)
        ]
        stats = {
            "source": str(self.source_path) if self.source_path else None,
            "rows": len(df),
            "columns": len(df.columns),
            "duplicate_rows": int(df.duplicated().sum()),
            "memory_mb": round(float(df.memory_usage(deep=deep_memory).sum() / 1024**2), 3),
            "load_warnings": list(self.load_warnings),
            "column_info": column_info,
        }
        self._cached_full_stats = stats
        self._cached_full_stats_version = self._df_version
        return stats

    def profile(self, sample_rows: int = 5) -> dict[str, Any]:
        """计算并返回数据集概况，包含 LRU 缓存避免重复计算。

        缓存策略：以 (df_version, sample_rows) 为键，同一 DataFrame 对象且
        sample_rows 相同时直接返回缓存。dataframe setter 会自动清除缓存。
        最多保留 _PROFILE_CACHE_MAX_ENTRIES 个条目，超过后按 LRU 策略
        淘汰最久未使用的条目（而非全清重建），提升缓存命中率。

        昂贵全表扫描（nunique、duplicated、memory_usage）通过
        ``_compute_full_stats`` 单独缓存，不同 sample_rows 的调用复用同一份。

        Args:
            sample_rows: 返回的样例行数（1-20）。

        Returns:
            包含 rows、columns、column_info、duplicate_rows、memory_mb、
            sample 等字段的字典。
        """
        sample_rows = max(1, min(int(sample_rows), 20))
        # 使用 (df_version, sample_rows) 作为缓存键，df_version 在每次 setter
        # 递增，避免 id(df) 复用（Python GC 后 id 可能重复）导致过期缓存。
        cache_key = self._df_version * 100 + sample_rows
        cached = self._profile_cache.get(cache_key)
        if cached is not None:
            # LRU：命中时移到末尾，标记为最近使用。
            self._profile_cache.move_to_end(cache_key)
            return cached
        # Reuse expensive full-table stats (nunique, duplicated, memory)
        # across different sample_rows calls.
        stats = self._compute_full_stats()
        result = to_jsonable({
            **stats,
            "sample": self.dataframe.head(sample_rows),
        })
        # LRU 淘汰：超容量时移除最旧条目（OrderedDict 首项），而非全清。
        # 全清会导致相邻两次不同 sample_rows 的调用互相淘汰，命中率归零。
        while len(self._profile_cache) >= _PROFILE_CACHE_MAX_ENTRIES:
            self._profile_cache.popitem(last=False)
        self._profile_cache[cache_key] = result
        return result

    @property
    def source_row_count(self) -> int:
        """Number of rows loaded from the uploaded source before agent mutations."""
        return self._source_row_count

    def save_dataframe(
        self,
        filename: str = "cleaned_data.csv",
        dataframe: pd.DataFrame | None = None,
        description: str = "清洗或变换后的数据集",
    ) -> Path:
        safe_name = re.sub(r"[^\w.\-\u4e00-\u9fff]", "_", Path(filename).name)
        path = (self.artifacts_dir / safe_name).resolve()
        frame = self.dataframe if dataframe is None else dataframe
        suffix = path.suffix.lower()
        if suffix == ".csv":
            frame.to_csv(path, index=False, encoding="utf-8-sig")
        elif suffix == ".xlsx":
            frame.to_excel(path, index=False)
        elif suffix == ".parquet":
            frame.to_parquet(path, index=False)
        else:
            raise ValueError("数据产物仅支持 .csv、.xlsx 或 .parquet。")
        self.register_artifact(path, "dataset", description)
        return path

    def ensure_plotly_bundle(self) -> Path | None:
        """Write the bundled Plotly.js once per workspace and reuse it for every chart.

        Returns the path to ``plotly.min.js`` inside this workspace's artifacts
        directory, or None when Plotly is not importable (extremely unlikely
        since it is a hard dependency). Each HTML artifact references this file
        relatively instead of inlining ~3 MB of JavaScript per chart.
        """
        try:
            from plotly.offline import get_plotlyjs
        except Exception:
            return None
        bundle = (self.artifacts_dir / PLOTLY_BUNDLE_NAME).resolve()
        if not bundle.exists():
            bundle.write_text(get_plotlyjs(), encoding="utf-8")
        return bundle

    def ensure_echarts_bundle(self) -> Path | None:
        """Write the bundled ECharts.js once per workspace and reuse it.

        首次调用时从官方 CDN 下载 ``echarts.min.js`` 到 artifacts_dir，
        后续复用。下载失败时返回 None，调用方需 fallback 到 CDN URL
        直接引用（在线场景可用，离线场景报错）。与 plotly bundle 同目录
        共存，互不冲突。
        """
        import urllib.request

        bundle = (self.artifacts_dir / ECHARTS_BUNDLE_NAME).resolve()
        if bundle.exists() and bundle.stat().st_size > 0:
            return bundle
        try:
            with urllib.request.urlopen(ECHARTS_CDN_URL, timeout=15) as response:  # noqa: S310
                content = response.read()
            if not content or len(content) < 1024:
                return None
            bundle.write_bytes(content)
            return bundle
        except Exception:
            # 离线 / 网络受限场景：返回 None，调用方走 CDN 直引 fallback。
            return None

    def ensure_echarts_gl_bundle(self) -> Path | None:
        """按需下载 echarts-gl 扩展 bundle，仅 3D 图表首次生成时触发。

        与 ``ensure_echarts_bundle`` 同策略：首次从 CDN 下载到 artifacts_dir
        后复用；失败返回 None，调用方 fallback 到 CDN URL 直引。
        """
        import urllib.request

        bundle = (self.artifacts_dir / ECHARTS_GL_BUNDLE_NAME).resolve()
        if bundle.exists() and bundle.stat().st_size > 0:
            return bundle
        try:
            with urllib.request.urlopen(ECHARTS_GL_CDN_URL, timeout=20) as response:  # noqa: S310
                content = response.read()
            if not content or len(content) < 1024:
                return None
            bundle.write_bytes(content)
            return bundle
        except Exception:
            return None

    def snapshot_state(self) -> tuple[pd.DataFrame, set[Path], int]:
        """Capture the active data and artifact files before one agent step.

        Returns the DataFrame (with copy-on-write-safe shallow copy), the set
        of existing artifact paths, and the current ``_df_version``.
        ``restore_state`` uses the version to short-circuit restoration when
        no DataFrame mutation occurred through the setter.
        """
        try:
            files = {
                path.resolve()
                for path in self.artifacts_dir.iterdir()
                if path.is_file()
            }
        except FileNotFoundError:
            files = set()
        # pandas >= 2.3 (enabled by default in 3.0) uses copy-on-write:
        # any modification to the live DataFrame triggers a CoW copy within
        # the modified column, so the shallow-copy reference remains intact.
        return self.dataframe.copy(deep=False), files, self._df_version

    def restore_state(self, snapshot: tuple[pd.DataFrame, set[Path], int]) -> None:
        """Rollback data mutations and files created by a failed agent step."""
        dataframe, existing_files, snapshot_version = snapshot
        if self._df_version != snapshot_version:
            self.dataframe = dataframe
        try:
            children = list(self.artifacts_dir.iterdir())
        except FileNotFoundError:
            children = []
        for path in children:
            resolved = path.resolve()
            if path.is_file() and resolved not in existing_files:
                path.unlink(missing_ok=True)
        self._artifacts = [item for item in self._artifacts if item.path.resolve() in existing_files]

    def save_checkpoint(self) -> Path:
        """Persist the active DataFrame separately from user-facing artifacts.

        原子写入：先写 .tmp 再 os.replace，避免进程被强杀（daemon 线程在进程
        退出时被杀）时留下半截 parquet 文件，导致下次 restore_checkpoint 读取
        到损坏文件抛 ArrowInvalid、会话无法恢复。
        """
        path = self.root / WORKSPACE_STATE_NAME
        temporary = path.with_suffix(".tmp")
        self.dataframe.to_parquet(temporary, index=False)
        os.replace(temporary, path)
        return path

    def restore_checkpoint(self) -> bool:
        """Restore the most recently persisted active DataFrame after a restart."""
        path = self.root / WORKSPACE_STATE_NAME
        if not path.is_file():
            return False
        self.dataframe = pd.read_parquet(path)
        return True

    def cleanup(self) -> None:
        """Remove this isolated workspace when a session expires or upload fails."""
        if self.root.exists():
            shutil.rmtree(self.root)

    def restore_artifacts(self, metadata: list[dict[str, str]] | None = None) -> None:
        """Re-register files already present when a persistent workspace is reopened."""
        if not self.artifacts_dir.is_dir():
            return
        metadata_by_name = {
            item.get("name", ""): item
            for item in (metadata or [])
            if isinstance(item, dict) and item.get("name")
        }
        for path in sorted(self.artifacts_dir.iterdir()):
            if not path.is_file() or path.name in (PLOTLY_BUNDLE_NAME, ECHARTS_BUNDLE_NAME):
                continue
            suffix = path.suffix.lower()
            saved = metadata_by_name.get(path.name, {})
            kind = saved.get("kind") or (
                "visualization"
                if suffix == ".html"
                else "image"
                if suffix in {".png", ".jpg", ".jpeg"}
                else "chart_data"
                if path.name.endswith(".plotly.json")
                else "dataset"
            )
            description = saved.get("description") or path.stem
            self.register_artifact(path, kind, description)
